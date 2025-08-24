# train.py

import torch
import torch.optim as optim
from torch.utils.data import TensorDataset, DataLoader
from sklearn.model_selection import StratifiedKFold
from sklearn.metrics import accuracy_score, f1_score, classification_report
import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches

class EarlyStopping:
    """Early stops the training if validation loss doesn't improve after a given patience."""
    def __init__(self, patience=7, verbose=False, delta=0, path='checkpoint.pt', trace_func=print):
        """
        Args:
            patience (int): How long to wait after last time validation loss improved.
            verbose (bool): If True, prints a message for each validation loss improvement. 
            delta (float): Minimum change in the monitored quantity to qualify as an improvement.
            path (str): Path for the checkpoint to be saved to.
            trace_func (function): trace print function.
        """
        self.patience = patience
        self.verbose = verbose
        self.counter = 0
        self.best_score = None
        self.early_stop = False
        # CORRECTED LINE: Changed np.Inf to np.inf for NumPy 2.0 compatibility
        self.val_loss_min = np.inf
        self.delta = delta
        self.path = path
        self.trace_func = trace_func

    def __call__(self, val_loss, model):
        score = -val_loss
        if self.best_score is None:
            self.best_score = score
            self.save_checkpoint(val_loss, model)
        elif score < self.best_score + self.delta:
            self.counter += 1
            self.trace_func(f'EarlyStopping counter: {self.counter} out of {self.patience}')
            if self.counter >= self.patience:
                self.early_stop = True
        else:
            self.best_score = score
            self.save_checkpoint(val_loss, model)
            self.counter = 0

    def save_checkpoint(self, val_loss, model):
        '''Saves model when validation loss decrease.'''
        if self.verbose:
            self.trace_func(f'Validation loss decreased ({self.val_loss_min:.6f} --> {val_loss:.6f}).  Saving model ...')
        torch.save(model.state_dict(), self.path)
        self.val_loss_min = val_loss

def train_model(model, train_loader, optimizer, criterion, device):
    """
    Performs one training epoch and returns the average loss.
    """
    model.train()
    total_loss = 0
    for data, target in train_loader:
        data, target = data.to(device), target.to(device)
        
        optimizer.zero_grad()
        output = model(data)
        loss = criterion(output, target)
        loss.backward()
        optimizer.step()
        
        total_loss += loss.item()
        
    return total_loss / len(train_loader)

def evaluate_model(model, test_loader, criterion, device):
    """
    Evaluates the model, returning accuracy, F1-score, a classification report, and average loss.
    """
    model.eval()
    all_preds = []
    all_targets = []
    total_loss = 0
    with torch.no_grad():
        for data, target in test_loader:
            data, target = data.to(device), target.to(device)
            output = model(data)
            
            loss = criterion(output, target)
            total_loss += loss.item()
            
            preds = torch.argmax(output, dim=1)
            
            all_preds.extend(preds.cpu().numpy())
            all_targets.extend(target.cpu().numpy())
            
    accuracy = accuracy_score(all_targets, all_preds)
    f1 = f1_score(all_targets, all_preds, average='weighted', zero_division=0)
    report = classification_report(all_targets, all_preds, zero_division=0)
    avg_loss = total_loss / len(test_loader)
    
    return accuracy, f1, report, avg_loss

def run_cross_validation(model_class, features, labels, subject_id, experiment_name, n_splits=5, epochs=50, batch_size=32, lr=0.001, patience=10):
    """
    Runs a full subject-dependent cross-validation loop with early stopping and generates a .docx report.
    """
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    skf = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=42)
    
    doc = Document()
    doc.add_heading(f'Results for Subject {subject_id} - Experiment: {experiment_name}', level=1)
    
    fold_accuracies, fold_f1_scores = [], []
    input_features_dim = features.shape[2]

    for fold, (train_idx, val_idx) in enumerate(skf.split(features, labels)):
        fold_header = f"--- Subject {subject_id} | Experiment: {experiment_name} | Fold {fold+1}/{n_splits} ---"
        print(fold_header)
        doc.add_heading(fold_header, level=2)
        
        X_train, X_val = features[train_idx], features[val_idx]
        y_train, y_val = labels[train_idx], labels[val_idx]
        
        val_counts = np.bincount(y_val)
        if len(val_counts) < 2 or np.min(val_counts) < 2:
            warning_msg = f"Warning: Skipping Fold {fold+1}. Validation set is not well-distributed: {val_counts}."
            print(warning_msg)
            doc.add_paragraph(warning_msg)
            continue

        X_train_tensor = torch.tensor(X_train, dtype=torch.float32)
        y_train_tensor = torch.tensor(y_train, dtype=torch.long)
        X_val_tensor = torch.tensor(X_val, dtype=torch.float32)
        y_val_tensor = torch.tensor(y_val, dtype=torch.long)
        
        train_dataset = TensorDataset(X_train_tensor, y_train_tensor)
        val_dataset = TensorDataset(X_val_tensor, y_val_tensor)
        train_loader = DataLoader(train_dataset, batch_size=batch_size, shuffle=True)
        val_loader = DataLoader(val_dataset, batch_size=batch_size, shuffle=False)

        if model_class.__name__ == 'BaselineCNN1D':
            model = model_class(input_features=input_features_dim).to(device)
        elif model_class.__name__ == 'RGGAT':
            model = model_class(initial_feature_dim=input_features_dim).to(device)
        else:
            raise ValueError("Unsupported model class")

        optimizer = optim.Adam(model.parameters(), lr=lr)
        criterion = torch.nn.CrossEntropyLoss()
        
        checkpoint_path = f'checkpoint_s{subject_id}_{experiment_name}_fold{fold+1}.pt'
        early_stopping = EarlyStopping(patience=patience, verbose=True, path=checkpoint_path)

        train_losses, val_losses = [], []
        for epoch in range(epochs):
            train_loss = train_model(model, train_loader, optimizer, criterion, device)
            _, _, _, val_loss = evaluate_model(model, val_loader, criterion, device)
            train_losses.append(train_loss)
            val_losses.append(val_loss)

            epoch_log = f"Epoch {epoch+1}/{epochs} | Train Loss: {train_loss:.4f} | Val Loss: {val_loss:.4f}"
            print(epoch_log)
            if (epoch + 1) % 10 == 0:
                doc.add_paragraph(epoch_log, style='List Bullet')
            
            early_stopping(val_loss, model)
            if early_stopping.early_stop:
                stop_msg = "Early stopping"
                print(stop_msg)
                doc.add_paragraph(stop_msg)
                break
        
        model.load_state_dict(torch.load(checkpoint_path))
        
        accuracy, f1, report, _ = evaluate_model(model, val_loader, criterion, device)
        
        doc.add_heading(f"\nFold {fold+1} Validation Results (from best model):", level=3)
        print(f"\nFold {fold+1} Validation Results (from best model):")
        
        results_summary = f"Accuracy: {accuracy:.4f}, F1-Score: {f1:.4f}"
        print(results_summary)
        doc.add_paragraph(results_summary)
        
        print("Classification Report:")
        print(report)
        doc.add_heading("Classification Report:", level=4)
        doc.add_paragraph(report)
        
        fold_accuracies.append(accuracy)
        fold_f1_scores.append(f1)
        
        plt.figure(figsize=(10, 6))
        plt.plot(train_losses, label='Training Loss')
        plt.plot(val_losses, label='Validation Loss')
        plt.axvline(x=len(train_losses)-1, color='r', linestyle='--', label='Early Stopping Point')
        plt.xlabel('Epochs'); plt.ylabel('Loss')
        plt.title(f'Loss Curve: Subj {subject_id}, Exp {experiment_name}, Fold {fold+1}')
        plt.legend(); plt.grid(True)
        
        results_dir = 'results'
        os.makedirs(results_dir, exist_ok=True)
        plot_filename = os.path.join(results_dir, f'loss_s{subject_id:02d}_{experiment_name}_fold{fold+1}.png')
        plt.savefig(plot_filename)
        plt.close()
        
        print(f"Saved loss curve to {plot_filename}\n")
        doc.add_picture(plot_filename, width=Inches(6.0))

    if not fold_accuracies:
        summary_msg = "\n--- No folds were successfully trained. Cannot compute summary. ---"
        print(summary_msg)
        doc.add_paragraph(summary_msg)
        avg_accuracy, avg_f1 = 0.0, 0.0
    else:
        avg_accuracy = np.mean(fold_accuracies)
        avg_f1 = np.mean(fold_f1_scores)
        
        summary_header = "\n--- Cross-Validation Summary ---"
        print(summary_header)
        doc.add_heading(summary_header, level=2)
        
        summary_text = f"Average Accuracy over {len(fold_accuracies)} folds: {avg_accuracy:.4f}\n"
        summary_text += f"Average F1-Score over {len(fold_f1_scores)} folds: {avg_f1:.4f}"
        print(summary_text)
        doc.add_paragraph(summary_text)
    
    report_filename = os.path.join(results_dir, f'report_s{subject_id:02d}_{experiment_name}.docx')
    doc.save(report_filename)
    print(f"\nSaved full report to {report_filename}")
    
    return avg_accuracy, avg_f1